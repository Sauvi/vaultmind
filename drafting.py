"""
drafting.py — VaultMind v0.9
Document drafting assistant.
Template-first approach — instant generation, optional AI customization.
Supports: NDA, Service Agreement, Employment Contract, Consulting Agreement
"""

from datetime import datetime
from pathlib import Path


# ══════════════════════════════════════════════════════════════
# TEMPLATES
# ══════════════════════════════════════════════════════════════

def _format_date(date_str: str) -> str:
    """Convert any date format to readable string like January 15, 2025."""
    import re
    if not date_str:
        return datetime.now().strftime("%B %d, %Y")
    # Already formatted like "January 15, 2025"
    if re.match(r"^[A-Za-z]+ \d{1,2}, \d{4}$", date_str):
        return date_str
    # HTML date input: YYYY-MM-DD
    if re.match(r"^\d{4}-\d{2}-\d{2}$", date_str):
        try:
            return datetime.strptime(date_str, "%Y-%m-%d").strftime("%B %d, %Y")
        except: pass
    # dd/mm/yyyy
    if re.match(r"^\d{1,2}/\d{1,2}/\d{4}$", date_str):
        try:
            return datetime.strptime(date_str, "%d/%m/%Y").strftime("%B %d, %Y")
        except: pass
    # YYYYMMDD
    if re.match(r"^\d{8}$", date_str):
        try:
            return datetime.strptime(date_str, "%Y%m%d").strftime("%B %d, %Y")
        except: pass
    # dd-mm-yyyy
    if re.match(r"^\d{1,2}-\d{1,2}-\d{4}$", date_str):
        try:
            return datetime.strptime(date_str, "%d-%m-%Y").strftime("%B %d, %Y")
        except: pass
    return date_str


TEMPLATES = {
    "nda": {
        "name": "Non-Disclosure Agreement (NDA)",
        "fields": ["party_a", "party_b", "effective_date", "duration_years",
                   "governing_law", "purpose"],
        "defaults": {
            "duration_years": "2",
            "governing_law":  "India",
            "purpose":        "evaluation of a potential business relationship",
        }
    },
    "service_agreement": {
        "name": "Service Agreement",
        "fields": ["client_name", "vendor_name", "effective_date",
                   "service_description", "total_amount", "currency",
                   "payment_terms", "governing_law"],
        "defaults": {
            "currency":      "INR",
            "payment_terms": "30 days from invoice",
            "governing_law": "India",
        }
    },
    "employment": {
        "name": "Employment Contract",
        "fields": ["employer_name", "employee_name", "effective_date",
                   "job_title", "department", "salary", "currency",
                   "notice_period_days", "governing_law"],
        "defaults": {
            "currency":           "INR",
            "notice_period_days": "30",
            "governing_law":      "India",
        }
    },
    "consulting": {
        "name": "Consulting Agreement",
        "fields": ["client_name", "consultant_name", "effective_date",
                   "scope_of_work", "daily_rate", "currency",
                   "duration_months", "governing_law"],
        "defaults": {
            "currency":        "INR",
            "duration_months": "6",
            "governing_law":   "India",
        }
    },
}


def get_template_list() -> list[dict]:
    return [
        {"id": k, "name": v["name"], "fields": v["fields"]}
        for k, v in TEMPLATES.items()
    ]


def get_template_fields(template_id: str) -> dict:
    if template_id not in TEMPLATES:
        raise ValueError(f"Unknown template: {template_id}")
    t = TEMPLATES[template_id]
    return {
        "id":       template_id,
        "name":     t["name"],
        "fields":   t["fields"],
        "defaults": t["defaults"],
    }


def draft_document(template_id: str, fields: dict) -> str:
    """Generate a draft document from template + fields."""
    if template_id == "nda":
        return _draft_nda(fields)
    elif template_id == "service_agreement":
        return _draft_service_agreement(fields)
    elif template_id == "employment":
        return _draft_employment(fields)
    elif template_id == "consulting":
        return _draft_consulting(fields)
    else:
        raise ValueError(f"Unknown template: {template_id}")


def save_draft(content: str, template_id: str,
               output_dir: str = "workspace/output") -> str:
    """Save draft as .txt file and return path."""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename  = f"draft_{template_id}_{timestamp}.txt"
    out_path  = Path(output_dir) / filename
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(content, encoding="utf-8")
    return str(out_path)


def save_draft_docx(content: str, template_id: str,
                    output_dir: str = "workspace/output") -> str:
    """Save draft as .docx file and return path."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError("Run: pip install python-docx")

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename  = f"draft_{template_id}_{timestamp}.docx"
    out_path  = Path(output_dir) / filename
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    for line in content.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph()
        elif line.startswith("===") or line.startswith("---"):
            continue
        elif line.isupper() and len(line) < 60:
            h = doc.add_heading(line.title(), level=2)
            h.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        else:
            p = doc.add_paragraph(line)
            if p.runs:
                p.runs[0].font.size = Pt(11)

    # Footer
    doc.add_paragraph()
    note = doc.add_paragraph(
        "DRAFT — Generated by VaultMind. Review by qualified legal counsel before use."
    )
    note.runs[0].font.size = Pt(9)
    note.runs[0].font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)
    note.runs[0].bold = True

    doc.save(str(out_path))
    return str(out_path)


# ══════════════════════════════════════════════════════════════
# TEMPLATE GENERATORS
# ══════════════════════════════════════════════════════════════

def _f(fields: dict, key: str, default: str = "[___]") -> str:
    """Get field value or placeholder."""
    return fields.get(key, "") or TEMPLATES.get(
        fields.get("_template_id", ""), {}
    ).get("defaults", {}).get(key, default)


def _draft_nda(f: dict) -> str:
    party_a      = f.get("party_a", "[Party A Name]")
    party_b      = f.get("party_b", "[Party B Name]")
    eff_date     = _format_date(f.get("effective_date", ""))
    duration     = f.get("duration_years", "2")
    gov_law      = f.get("governing_law", "India")
    purpose      = f.get("purpose", "evaluation of a potential business relationship")

    return f"""NON-DISCLOSURE AGREEMENT
================================================================================

This Non-Disclosure Agreement ("Agreement") is entered into as of {eff_date}
("Effective Date") by and between:

{party_a} (hereinafter referred to as "Disclosing Party"), AND
{party_b} (hereinafter referred to as "Receiving Party").

RECITALS

The parties wish to explore {purpose} and may disclose certain confidential
and proprietary information to each other. This Agreement sets forth the terms
under which such information will be protected.

ARTICLE 1 — DEFINITIONS

1.1 "Confidential Information" means any and all non-public information
disclosed by the Disclosing Party to the Receiving Party, whether orally,
in writing, or by any other means, that is designated as confidential or
that reasonably should be understood to be confidential.

1.2 Confidential Information does not include information that:
(a) is or becomes publicly available through no breach of this Agreement;
(b) was rightfully known to the Receiving Party before disclosure;
(c) is independently developed by the Receiving Party without use of
    Confidential Information; or
(d) is required to be disclosed by applicable law or court order.

ARTICLE 2 — CONFIDENTIALITY OBLIGATIONS

2.1 The Receiving Party agrees to:
(a) hold all Confidential Information in strict confidence;
(b) not disclose Confidential Information to any third party without prior
    written consent of the Disclosing Party;
(c) use Confidential Information solely for the purpose of {purpose};
(d) protect Confidential Information using at least the same degree of care
    it uses to protect its own confidential information, but no less than
    reasonable care.

2.2 The Receiving Party may disclose Confidential Information only to its
employees, officers, and advisors who have a need to know and are bound by
confidentiality obligations no less restrictive than this Agreement.

ARTICLE 3 — TERM

3.1 This Agreement shall commence on the Effective Date and remain in effect
for a period of {duration} years, unless earlier terminated by mutual written
agreement of the parties.

3.2 The confidentiality obligations shall survive termination of this Agreement
for an additional period of {duration} years with respect to all Confidential
Information disclosed during the term.

ARTICLE 4 — OWNERSHIP

4.1 All Confidential Information remains the exclusive property of the
Disclosing Party. Nothing in this Agreement grants the Receiving Party any
rights in or to the Confidential Information except as expressly set forth.

ARTICLE 5 — REMEDIES

5.1 The parties acknowledge that any breach of this Agreement may cause
irreparable harm for which monetary damages would be inadequate. Each party
shall be entitled to seek injunctive or other equitable relief in addition to
all other remedies available at law.

ARTICLE 6 — GOVERNING LAW

6.1 This Agreement shall be governed by and construed in accordance with the
laws of {gov_law}, without regard to its conflict of law provisions.

6.2 Any disputes arising under this Agreement shall be subject to the
exclusive jurisdiction of the competent courts in {gov_law}.

ARTICLE 7 — GENERAL

7.1 This Agreement constitutes the entire agreement between the parties
concerning its subject matter and supersedes all prior agreements.

7.2 This Agreement may be amended only by a written instrument signed by
both parties.

7.3 If any provision is found invalid or unenforceable, the remaining
provisions shall continue in full force.

================================================================================
SIGNATURES

{party_a}

Signature: _______________________
Name:      ________________________
Title:     ________________________
Date:      ________________________


{party_b}

Signature: _______________________
Name:      ________________________
Title:     ________________________
Date:      ________________________

================================================================================
DRAFT — Generated by VaultMind. Review by qualified legal counsel before use.
================================================================================"""


def _draft_service_agreement(f: dict) -> str:
    client   = f.get("client_name",        "[Client Name]")
    vendor   = f.get("vendor_name",        "[Vendor Name]")
    eff_date = _format_date(f.get("effective_date", ""))
    services = f.get("service_description","[Description of services to be provided]")
    amount   = f.get("total_amount",       "[Amount]")
    currency = f.get("currency",           "INR")
    payment  = f.get("payment_terms",      "30 days from invoice")
    gov_law  = f.get("governing_law",      "India")

    return f"""SERVICE AGREEMENT
================================================================================

This Service Agreement ("Agreement") is entered into as of {eff_date}
("Effective Date") by and between:

{client} (hereinafter referred to as "Client"), AND
{vendor} (hereinafter referred to as "Service Provider").

ARTICLE 1 — SERVICES

1.1 Service Provider shall provide the following services ("Services"):
{services}

1.2 Service Provider shall perform the Services in a professional and
workmanlike manner using qualified personnel.

1.3 Any changes to the scope of Services must be agreed in writing by
both parties prior to implementation.

ARTICLE 2 — COMPENSATION

2.1 In consideration for the Services, Client shall pay Service Provider
{currency} {amount} as per the payment schedule mutually agreed.

2.2 Service Provider shall submit invoices to Client, and Client shall make
payment within {payment} of receiving a valid invoice.

2.3 Late payments shall accrue interest at 18% per annum from the due date.

ARTICLE 3 — INTELLECTUAL PROPERTY

3.1 All work product and deliverables created by Service Provider specifically
for Client under this Agreement ("Deliverables") shall be owned by Client.

3.2 Service Provider retains ownership of all pre-existing tools, frameworks,
and methodologies ("Background IP"). A non-exclusive license to use Background
IP as part of the Deliverables is granted to Client.

ARTICLE 4 — CONFIDENTIALITY

4.1 Both parties agree to maintain the confidentiality of each other's
proprietary and confidential information disclosed in connection with this
Agreement for a period of 3 years after termination.

ARTICLE 5 — TERM AND TERMINATION

5.1 This Agreement commences on the Effective Date and continues until
completion of the Services, unless earlier terminated.

5.2 Either party may terminate this Agreement with 30 days' written notice.

5.3 Client may terminate immediately for cause if Service Provider materially
breaches this Agreement and fails to cure within 15 days of written notice.

ARTICLE 6 — LIMITATION OF LIABILITY

6.1 Neither party shall be liable for indirect, incidental, or consequential
damages arising from this Agreement.

6.2 Each party's total liability shall not exceed the total fees paid or
payable under this Agreement in the preceding 12 months.

ARTICLE 7 — GOVERNING LAW

7.1 This Agreement shall be governed by the laws of {gov_law}.

================================================================================
SIGNATURES

{client}                          {vendor}

Signature: ___________________    Signature: ___________________
Name:      ___________________    Name:      ___________________
Title:     ___________________    Title:     ___________________
Date:      ___________________    Date:      ___________________

================================================================================
DRAFT — Generated by VaultMind. Review by qualified legal counsel before use.
================================================================================"""


def _draft_employment(f: dict) -> str:
    employer = f.get("employer_name",      "[Employer Name]")
    employee = f.get("employee_name",      "[Employee Name]")
    eff_date = _format_date(f.get("effective_date", ""))
    title    = f.get("job_title",          "[Job Title]")
    dept     = f.get("department",         "[Department]")
    salary   = f.get("salary",             "[Salary Amount]")
    currency = f.get("currency",           "INR")
    notice   = f.get("notice_period_days", "30")
    gov_law  = f.get("governing_law",      "India")

    return f"""EMPLOYMENT CONTRACT
================================================================================

This Employment Contract ("Agreement") is entered into as of {eff_date}
by and between:

{employer} (hereinafter referred to as "Employer"), AND
{employee} (hereinafter referred to as "Employee").

ARTICLE 1 — EMPLOYMENT

1.1 Employer hereby employs Employee in the capacity of {title} in the
{dept} department, effective {eff_date}.

1.2 Employee shall report to [Reporting Manager] and perform all duties
reasonably assigned from time to time.

1.3 Employee shall devote full working time and best efforts exclusively
to the Employer's business during normal working hours.

ARTICLE 2 — COMPENSATION

2.1 Employer shall pay Employee a gross salary of {currency} {salary}
per annum, payable monthly in arrears.

2.2 Salary shall be subject to applicable statutory deductions including
income tax (TDS), Provident Fund (PF), and Employee State Insurance (ESI)
as applicable.

2.3 Salary shall be reviewed annually at Employer's discretion based on
performance and business conditions.

ARTICLE 3 — WORKING HOURS AND LEAVE

3.1 Standard working hours are Monday to Friday, 9:00 AM to 6:00 PM,
with one hour for lunch, subject to operational requirements.

3.2 Employee is entitled to leave as per Employer's leave policy, including
annual leave, sick leave, and public holidays.

ARTICLE 4 — CONFIDENTIALITY

4.1 Employee shall keep confidential all trade secrets, client information,
business strategies, and proprietary information of Employer during and after
employment.

4.2 This obligation shall survive termination of employment for 3 years.

ARTICLE 5 — INTELLECTUAL PROPERTY

5.1 All inventions, works, and developments created by Employee in the course
of employment shall be the exclusive property of Employer.

ARTICLE 6 — TERMINATION

6.1 Either party may terminate this Agreement by giving {notice} days'
written notice to the other party.

6.2 Employer may terminate immediately without notice for cause, including
gross misconduct, breach of confidentiality, or fraud.

6.3 Upon termination, Employee shall return all company property, documents,
and data immediately.

ARTICLE 7 — GOVERNING LAW

7.1 This Agreement is governed by the laws of {gov_law} and the applicable
labour legislation.

================================================================================
SIGNATURES

{employer}                        {employee}

Signature: ___________________    Signature: ___________________
Name:      ___________________    Name:      ___________________
Title:     ___________________    Date:      ___________________
Date:      ___________________

================================================================================
DRAFT — Generated by VaultMind. Review by qualified legal counsel before use.
================================================================================"""


def _draft_consulting(f: dict) -> str:
    client    = f.get("client_name",     "[Client Name]")
    consult   = f.get("consultant_name", "[Consultant Name]")
    eff_date  = _format_date(f.get("effective_date", ""))
    scope     = f.get("scope_of_work",   "[Description of consulting work]")
    rate      = f.get("daily_rate",      "[Daily Rate]")
    currency  = f.get("currency",        "INR")
    duration  = f.get("duration_months", "6")
    gov_law   = f.get("governing_law",   "India")

    return f"""CONSULTING AGREEMENT
================================================================================

This Consulting Agreement ("Agreement") is entered into as of {eff_date}
("Effective Date") by and between:

{client} (hereinafter referred to as "Client"), AND
{consult} (hereinafter referred to as "Consultant").

ARTICLE 1 — CONSULTING SERVICES

1.1 Consultant agrees to provide the following consulting services:
{scope}

1.2 Consultant shall perform the services as an independent contractor.
Nothing in this Agreement creates an employer-employee relationship.

1.3 Consultant shall have the freedom to determine the method, details,
and means of performing the services, subject to Client's overall direction.

ARTICLE 2 — TERM

2.1 This Agreement commences on {eff_date} and continues for {duration}
months, unless earlier terminated or extended by mutual written agreement.

ARTICLE 3 — FEES

3.1 Client shall pay Consultant {currency} {rate} per day worked.

3.2 Consultant shall submit monthly invoices with timesheets, and Client
shall pay within 15 business days of receipt.

3.3 Client shall reimburse pre-approved reasonable expenses with receipts.

ARTICLE 4 — INDEPENDENT CONTRACTOR

4.1 Consultant is an independent contractor. Consultant is responsible for
all taxes on fees received and shall not be entitled to employee benefits.

4.2 Consultant may engage other clients during the term, provided there is
no conflict of interest with Client's business.

ARTICLE 5 — CONFIDENTIALITY

5.1 Consultant shall keep all Client information strictly confidential during
and for 3 years after the term of this Agreement.

ARTICLE 6 — INTELLECTUAL PROPERTY

6.1 All work product created specifically for Client under this Agreement
shall be owned by Client upon full payment of fees.

6.2 Consultant retains ownership of pre-existing methodologies, tools, and
frameworks, with a license granted to Client for agreed deliverables.

ARTICLE 7 — TERMINATION

7.1 Either party may terminate this Agreement with 14 days' written notice.

7.2 Client may terminate immediately if Consultant commits a material breach
that remains uncured for 7 days after written notice.

ARTICLE 8 — GOVERNING LAW

8.1 This Agreement is governed by the laws of {gov_law}.

================================================================================
SIGNATURES

{client}                          {consult}

Signature: ___________________    Signature: ___________________
Name:      ___________________    Name:      ___________________
Title:     ___________________    Date:      ___________________
Date:      ___________________

================================================================================
DRAFT — Generated by VaultMind. Review by qualified legal counsel before use.
================================================================================"""
