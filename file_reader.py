"""
file_reader.py — VaultMind v1.2
Smart document ingestion pipeline with OCR support.

NEW in v1.2:
  - OCR for scanned PDFs (Tesseract via pytesseract)
  - Hindi + English OCR (lang='hin+eng')
  - Auto-detection: digital PDF → direct extract, scanned → OCR
  - No code changes needed — OCR is transparent to the rest of the app

OCR requirements (Windows):
  1. pip install pytesseract pillow pdf2image
  2. Install Tesseract: https://github.com/UB-Mannheim/tesseract/wiki
     Default path: C:\\Program Files\\Tesseract-OCR\\tesseract.exe
  3. Install Poppler: https://github.com/oschwartz10612/poppler-windows/releases
     Add bin/ folder to PATH (needed by pdf2image)

After install, scanned PDFs will work automatically.

Supports: .txt, .md, .pdf (digital + scanned), .docx
"""

from pathlib import Path
from typing import NamedTuple
import re

SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx", ".md"}

# Hard cap for read_document (used for quick preview)
MAX_CHARS = 3000

# Tesseract Windows paths — auto-detected
_TESSERACT_PATHS = [
    r"C:\Program Files\Tesseract-OCR\tesseract.exe",
    r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    r"C:\Users\Public\Tesseract-OCR\tesseract.exe",
]

# Poppler path — read from installer config file if present
def _get_poppler_path() -> str | None:
    """Read Poppler path saved by install.bat, or search common locations."""
    # Check installer-written config first
    config = Path(__file__).parent / "poppler_path.txt"
    if config.exists():
        p = config.read_text(encoding="utf-8").strip()
        if p and Path(p).exists():
            return p
    # Check tools subfolder (installer puts it here)
    for tools_path in (Path(__file__).parent / "tools").glob("poppler*/Library/bin"):
        if tools_path.exists():
            return str(tools_path)
    for tools_path in (Path(__file__).parent / "tools").glob("poppler*/bin"):
        if tools_path.exists():
            return str(tools_path)
    # Let pdf2image find it via PATH
    return None

_POPPLER_PATH = _get_poppler_path()

# OCR language — English + Hindi. Add more as needed.
# Full list: https://tesseract-ocr.github.io/tessdoc/Data-Files-in-different-versions.html
OCR_LANG = "eng+hin"

# Page text density threshold — pages with fewer chars than this are treated as scanned
_SCANNED_PAGE_THRESHOLD = 50


class DocumentResult(NamedTuple):
    text:        str    # clean, normalised text ready for AI
    raw_chars:   int    # original character count before cleaning
    final_chars: int    # character count after cleaning + truncation
    pages:       int    # page count
    file_type:   str    # extension without dot
    truncated:   bool   # whether document was truncated
    is_ocr:      bool   # whether OCR was used


# ── OCR helpers ───────────────────────────────────────────────────────────────

def _configure_tesseract() -> bool:
    """Set Tesseract path on Windows. Returns True if found."""
    try:
        import pytesseract
        # Try auto-detect first
        import shutil
        if shutil.which("tesseract"):
            return True  # Already on PATH
        # Try known Windows paths
        for p in _TESSERACT_PATHS:
            if Path(p).exists():
                pytesseract.pytesseract.tesseract_cmd = p
                return True
        return False
    except ImportError:
        return False


def _is_ocr_available() -> bool:
    """Check if both pytesseract and pdf2image are available."""
    try:
        import pytesseract
        import pdf2image
        return _configure_tesseract()
    except ImportError:
        return False


def _ocr_page_image(image) -> str:
    """Run Tesseract OCR on a PIL Image. Returns extracted text."""
    import pytesseract
    _configure_tesseract()
    try:
        text = pytesseract.image_to_string(
            image,
            lang=OCR_LANG,
            config="--psm 1 --oem 3",  # PSM 1 = auto page seg, OEM 3 = LSTM+legacy
        )
        return text.strip()
    except Exception as e:
        print(f"[OCR] Page OCR failed: {e}")
        return ""


def _extract_pdf_with_ocr(path: Path, full: bool = False) -> tuple[str, int, bool]:
    """
    Smart PDF extraction:
    1. Try digital text extraction first (PyMuPDF)
    2. For each page with < threshold chars → run OCR on that page
    3. Mix digital + OCR pages together

    Returns (text, page_count, used_ocr)
    """
    try:
        import fitz
    except ImportError:
        raise RuntimeError("Run: pip install pymupdf")

    size_mb = path.stat().st_size / (1024 * 1024)
    max_size = 50 if full else 15
    if size_mb > max_size:
        raise RuntimeError(
            f"PDF is {size_mb:.1f} MB. Max is {max_size} MB. "
            "Split the document or convert to DOCX."
        )

    ocr_available  = _is_ocr_available()
    used_ocr       = False
    pages_text     = []
    max_chars_soft = None if full else MAX_CHARS * 3

    with fitz.open(str(path)) as doc:
        total_pages = len(doc)

        for page_num in range(total_pages):
            page = doc[page_num]

            # ── Try digital extraction first ──────────────────────────────
            blocks = page.get_text("blocks")
            if blocks:
                blocks_sorted = sorted(blocks, key=lambda b: (round(b[1] / 20) * 20, b[0]))
                page_text     = "\n".join(
                    b[4].strip() for b in blocks_sorted
                    if isinstance(b[4], str) and b[4].strip()
                )
            else:
                page_text = page.get_text("text")

            # ── OCR fallback for scanned pages ────────────────────────────
            if len(page_text.strip()) < _SCANNED_PAGE_THRESHOLD and ocr_available:
                try:
                    from pdf2image import convert_from_path
                    # Convert just this page to image
                    images = convert_from_path(
                        str(path),
                        first_page=page_num + 1,
                        last_page=page_num + 1,
                        dpi=200,  # 200 DPI — good balance of speed vs accuracy
                        poppler_path=_POPPLER_PATH,  # auto-detected by installer
                    )
                    if images:
                        ocr_text = _ocr_page_image(images[0])
                        if len(ocr_text) > len(page_text):
                            page_text = ocr_text
                            used_ocr  = True
                            print(f"[OCR] Page {page_num + 1}: OCR extracted {len(ocr_text)} chars")
                except Exception as e:
                    print(f"[OCR] Page {page_num + 1} OCR failed: {e}")

            if page_text.strip():
                pages_text.append(f"[Page {page_num + 1}]\n{page_text.strip()}")

            # Early exit for non-full reads
            if max_chars_soft and sum(len(p) for p in pages_text) > max_chars_soft:
                if page_num < total_pages - 1:
                    pages_text.append(
                        f"[Pages {page_num + 2}–{total_pages} not read — "
                        f"sufficient content extracted]"
                    )
                break

    if not pages_text:
        if not ocr_available:
            raise RuntimeError(
                f"No text extracted from {path.name}. "
                "This appears to be a scanned PDF. "
                "To enable OCR, install: pip install pytesseract pillow pdf2image "
                "and Tesseract (https://github.com/UB-Mannheim/tesseract/wiki)"
            )
        else:
            raise RuntimeError(
                f"No text extracted from {path.name} even with OCR. "
                "The PDF may be corrupt or in an unsupported format."
            )

    return "\n\n".join(pages_text), total_pages, used_ocr


# ── Public API ────────────────────────────────────────────────────────────────

def read_file(file_path: str) -> str:
    """Legacy API — returns clean text string. Used by actions.py."""
    return read_document(file_path).text


def read_document(file_path: str) -> DocumentResult:
    """Full pipeline — returns DocumentResult with text + metadata."""
    path = Path(file_path)
    ext  = path.suffix.lower()

    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported format: '{ext}'. "
            f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    is_ocr = False

    if ext in (".txt", ".md"):
        raw, pages = _extract_text(path)
    elif ext == ".pdf":
        raw, pages, is_ocr = _extract_pdf_with_ocr(path, full=False)
    elif ext == ".docx":
        raw, pages = _extract_docx(path)

    raw_chars = len(raw)
    cleaned   = _clean(raw)
    final, truncated = _smart_truncate(cleaned, MAX_CHARS)

    return DocumentResult(
        text        = final,
        raw_chars   = raw_chars,
        final_chars = len(final),
        pages       = pages,
        file_type   = ext.lstrip("."),
        truncated   = truncated,
        is_ocr      = is_ocr,
    )


def read_full_text(file_path: str) -> str:
    """
    Extract complete text without truncation.
    Used by RAG indexing — we need the full document.
    Returns (text, is_ocr).
    """
    path = Path(file_path)
    ext  = path.suffix.lower()

    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported format: '{ext}'")

    if ext in (".txt", ".md"):
        raw, _ = _extract_text(path)
        return _clean(raw)
    elif ext == ".pdf":
        raw, _, _ = _extract_pdf_with_ocr(path, full=True)
        return _clean(raw)
    elif ext == ".docx":
        raw, _ = _extract_docx(path)
        return _clean(raw)

    return ""


def read_full_text_with_meta(file_path: str) -> tuple[str, int, bool]:
    """
    Full extraction returning (text, pages, is_ocr).
    Used by app.py on upload for RAG indexing.
    """
    path = Path(file_path)
    ext  = path.suffix.lower()

    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported format: '{ext}'")

    if ext in (".txt", ".md"):
        raw, pages = _extract_text(path)
        return _clean(raw), pages, False
    elif ext == ".pdf":
        raw, pages, is_ocr = _extract_pdf_with_ocr(path, full=True)
        return _clean(raw), pages, is_ocr
    elif ext == ".docx":
        raw, pages = _extract_docx(path)
        return _clean(raw), pages, False

    return "", 0, False


def get_file_info(file_path: str) -> dict:
    path   = Path(file_path)
    ext    = path.suffix.lower()
    size_kb = round(path.stat().st_size / 1024, 1) if path.exists() else 0
    return {
        "name":      path.name,
        "extension": ext,
        "size_kb":   size_kb,
        "supported": ext in SUPPORTED_EXTENSIONS,
    }


def ocr_status() -> dict:
    """Return OCR availability status for display in UI."""
    available     = _is_ocr_available()
    tesseract_ok  = False
    pdf2image_ok  = False
    pytesseract_ok = False

    try:
        import pytesseract
        pytesseract_ok = True
        tesseract_ok   = _configure_tesseract()
    except ImportError:
        pass

    try:
        import pdf2image
        pdf2image_ok = True
    except ImportError:
        pass

    return {
        "available":      available,
        "pytesseract":    pytesseract_ok,
        "pdf2image":      pdf2image_ok,
        "tesseract_bin":  tesseract_ok,
        "ocr_lang":       OCR_LANG,
        "message": (
            "OCR ready — Hindi + English scanned PDFs supported"
            if available else
            "OCR not available. Install: pip install pytesseract pillow pdf2image "
            "+ Tesseract from https://github.com/UB-Mannheim/tesseract/wiki"
        )
    }


# ── Extractors ────────────────────────────────────────────────────────────────

def _extract_text(path: Path) -> tuple[str, int]:
    """Plain text / markdown — try encodings gracefully."""
    for enc in ["utf-8", "utf-8-sig", "latin-1", "cp1252"]:
        try:
            text  = path.read_text(encoding=enc)
            lines = len([l for l in text.splitlines() if l.strip()])
            return text, max(1, lines // 40)
        except UnicodeDecodeError:
            continue
    raise RuntimeError(f"Cannot decode {path.name}. Save as UTF-8 and retry.")


def _extract_docx(path: Path) -> tuple[str, int]:
    """DOCX extraction preserving document structure."""
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("Run: pip install python-docx")

    doc   = Document(str(path))
    parts = []

    for para in doc.paragraphs:
        text  = para.text.strip()
        if not text:
            continue
        style = para.style.name.lower() if para.style else ""
        if "heading 1" in style:
            parts.append(f"\n## {text}")
        elif "heading 2" in style:
            parts.append(f"\n### {text}")
        elif "heading" in style:
            parts.append(f"\n#### {text}")
        elif "list" in style or text.startswith(("•", "-", "*", "◦")):
            parts.append(f" - {text.lstrip('•-*◦ ')}")
        else:
            parts.append(text)

    for table in doc.tables:
        table_rows = []
        for i, row in enumerate(table.rows):
            cells  = [c.text.strip().replace("\n", " ") for c in row.cells]
            seen   = []
            deduped = []
            for c in cells:
                if c not in seen:
                    deduped.append(c)
                    seen.append(c)
                else:
                    deduped.append("")
            row_str = " | ".join(deduped)
            table_rows.append(row_str)
            if i == 0:
                table_rows.append("-" * len(row_str))
        if table_rows:
            parts.append("\n" + "\n".join(table_rows))

    if not parts:
        raise RuntimeError(f"No text found in {path.name}.")

    text       = "\n".join(parts)
    word_count = len(text.split())
    pages      = max(1, word_count // 300)
    return text, pages


# ── Cleaner ───────────────────────────────────────────────────────────────────

def _clean(text: str) -> str:
    """
    Remove structural noise while preserving ALL language characters.
    Safe for Hindi/Devanagari, Arabic, Chinese, etc.
    """
    # Page number lines
    text = re.sub(r'(?m)^\s*[Pp]age\s+\d+\s+(?:of\s+\d+)?\s*$', '', text)
    text = re.sub(r'(?m)^\s*\d+\s*$', '', text)
    # Divider lines
    text = re.sub(r'[-_=*#~]{4,}', '', text)
    # Email headers
    text = re.sub(r'(?m)^(From|To|CC|BCC|Subject|Date)\s*:.*$', '', text)
    # Collapse excess newlines
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Collapse multiple spaces/tabs
    text = re.sub(r'[ \t]{2,}', ' ', text)
    # Strip pure-noise lines
    lines = []
    for line in text.splitlines():
        stripped = line.strip()
        if len(stripped) >= 3 or stripped.startswith('#'):
            lines.append(line)
        elif not stripped:
            lines.append('')
    text = '\n'.join(lines)
    # Remove ASCII control chars only (preserve all Unicode)
    text = ''.join(ch for ch in text if ord(ch) >= 32 or ch in '\n\r\t')
    return text.strip()


# ── Smart truncation ──────────────────────────────────────────────────────────

def _smart_truncate(text: str, max_chars: int) -> tuple[str, bool]:
    """
    Keep first 65% + last 35% for contracts
    (parties at start, signatures/dates at end).
    """
    if len(text) <= max_chars:
        return text, False

    head = int(max_chars * 0.65)
    tail = max_chars - head
    truncated = (
        text[:head]
        + "\n\n[... middle section omitted — showing start and end of document ...]\n\n"
        + text[-tail:]
    )
    return truncated, True


def clean_text(text: str) -> str:
    """Legacy alias — used by actions.py directly."""
    return _clean(text)
