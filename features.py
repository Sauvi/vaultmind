"""
features.py — VaultMind v0.8
Four new sellable capabilities:
1. contract_compare   — diff two contracts, highlight changes + risks
2. clause_extract     — pull specific clause types from any contract
3. deadline_tracker   — extract all dates + obligations as checklist
4. report_generator   — generate professional DOCX report
"""

import re
from datetime import datetime
from pathlib import Path
from legal_extractor import (
    extract_fast, _extract_dates, _extract_risk_sentences,
    _extract_obligation_sentences, _extract_parties,
    _detect_governing_law, _detect_contract_type
)


# ══════════════════════════════════════════════════════════════
# 1. CONTRACT COMPARISON
# ══════════════════════════════════════════════════════════════

def compare_contracts(text_a: str, text_b: str,
                      name_a: str = "Contract A",
                      name_b: str = "Contract B") -> dict:
    """
    Compare two contracts and return structured differences.
    Uses pattern extraction — no AI needed, instant.
    """
    ext_a = extract_fast(text_a)
    ext_b = extract_fast(text_b)

    def _set(lst): return set(s.strip().lower() for s in lst)

    # Party changes
    parties_added    = list(_set(ext_b.parties) - _set(ext_a.parties))
    parties_removed  = list(_set(ext_a.parties) - _set(ext_b.parties))

    # Date changes
    dates_added      = list(_set(ext_b.dates) - _set(ext_a.dates))
    dates_removed    = list(_set(ext_a.dates) - _set(ext_b.dates))

    # Risk clause changes
    risks_added      = []
    risks_removed    = []
    for r in ext_b.risk_clauses:
        if not any(_similarity(r, ra) > 0.6 for ra in ext_a.risk_clauses):
            risks_added.append(r[:200])
    for r in ext_a.risk_clauses:
        if not any(_similarity(r, rb) > 0.6 for rb in ext_b.risk_clauses):
            risks_removed.append(r[:200])

    # Obligation changes
    obls_added   = []
    obls_removed = []
    for o in ext_b.obligations:
        if not any(_similarity(o, oa) > 0.6 for oa in ext_a.obligations):
            obls_added.append(o[:200])
    for o in ext_a.obligations:
        if not any(_similarity(o, ob) > 0.6 for ob in ext_b.obligations):
            obls_removed.append(o[:200])

    # Governing law change
    law_changed = (ext_a.governing_law or "").strip() != (ext_b.governing_law or "").strip()

    # Contract type change
    type_changed = ext_a.contract_type != ext_b.contract_type

    # Risk score — higher = more different/riskier
    risk_score = (
        len(risks_added) * 3 +
        len(obls_added) * 2 +
        len(dates_added) * 1 +
        (5 if law_changed else 0) +
        (3 if type_changed else 0)
    )
    risk_level = "HIGH" if risk_score >= 10 else "MEDIUM" if risk_score >= 4 else "LOW"

    return {
        "name_a":          name_a,
        "name_b":          name_b,
        "contract_type_a": ext_a.contract_type,
        "contract_type_b": ext_b.contract_type,
        "type_changed":    type_changed,
        "governing_law_a": ext_a.governing_law,
        "governing_law_b": ext_b.governing_law,
        "law_changed":     law_changed,
        "parties_added":   parties_added,
        "parties_removed": parties_removed,
        "dates_added":     dates_added,
        "dates_removed":   dates_removed,
        "risks_added":     risks_added[:5],
        "risks_removed":   risks_removed[:5],
        "obls_added":      obls_added[:5],
        "obls_removed":    obls_removed[:5],
        "risk_score":      risk_score,
        "risk_level":      risk_level,
        "summary": _compare_summary(
            risk_level, type_changed, law_changed,
            len(risks_added), len(obls_added), len(dates_added)
        ),
    }


def _similarity(a: str, b: str) -> float:
    """Simple word-overlap similarity score."""
    wa = set(a.lower().split())
    wb = set(b.lower().split())
    if not wa or not wb:
        return 0.0
    return len(wa & wb) / max(len(wa), len(wb))


def _compare_summary(risk_level, type_changed, law_changed,
                     n_risks, n_obls, n_dates) -> str:
    parts = []
    if type_changed:
        parts.append("contract type changed")
    if law_changed:
        parts.append("governing law changed — significant legal implications")
    if n_risks > 0:
        parts.append(f"{n_risks} new risk clause(s) added")
    if n_obls > 0:
        parts.append(f"{n_obls} new obligation(s) added")
    if n_dates > 0:
        parts.append(f"{n_dates} new date(s) added")
    if not parts:
        return "Contracts are substantially similar — no major differences detected."
    return f"Risk level: {risk_level}. Key changes: {'; '.join(parts)}."


def format_comparison(result: dict) -> str:
    """Format comparison result as readable text."""
    lines = []
    lines.append("=" * 44)
    lines.append("CONTRACT COMPARISON REPORT")
    lines.append("=" * 44)
    lines.append(f"Document A : {result['name_a']}")
    lines.append(f"Document B : {result['name_b']}")
    lines.append(f"Risk level : {result['risk_level']}")
    lines.append("")
    lines.append(result["summary"])
    lines.append("")

    if result["type_changed"]:
        lines.append("⚠  CONTRACT TYPE CHANGED")
        lines.append(f"   A: {result['contract_type_a']}")
        lines.append(f"   B: {result['contract_type_b']}")
        lines.append("")

    if result["law_changed"]:
        lines.append("⚠  GOVERNING LAW CHANGED")
        lines.append(f"   A: {result['governing_law_a'] or 'Not specified'}")
        lines.append(f"   B: {result['governing_law_b'] or 'Not specified'}")
        lines.append("")

    if result["risks_added"]:
        lines.append(f"NEW RISKS IN {result['name_b']} ({len(result['risks_added'])})")
        for r in result["risks_added"]:
            lines.append(f"  ⚠  {r}")
        lines.append("")

    if result["risks_removed"]:
        lines.append(f"RISKS REMOVED FROM {result['name_a']} ({len(result['risks_removed'])})")
        for r in result["risks_removed"]:
            lines.append(f"  ✓  {r}")
        lines.append("")

    if result["obls_added"]:
        lines.append(f"NEW OBLIGATIONS IN {result['name_b']} ({len(result['obls_added'])})")
        for o in result["obls_added"]:
            lines.append(f"  →  {o}")
        lines.append("")

    if result["dates_added"]:
        lines.append(f"NEW DATES IN {result['name_b']}")
        for d in result["dates_added"]:
            lines.append(f"  📅  {d}")
        lines.append("")

    lines.append("=" * 44)
    return "\n".join(lines)


# ══════════════════════════════════════════════════════════════
# 2. CLAUSE EXTRACTION
# ══════════════════════════════════════════════════════════════

CLAUSE_TYPES = {
    "payment":         ["payment", "invoice", "fee", "price", "cost", "compensation", "salary", "remuneration", "inr", "usd", "amount"],
    "termination":     ["terminat", "cancel", "end", "expir", "cessation", "wind-down", "dissolv"],
    "confidentiality": ["confidential", "non-disclosure", "nda", "proprietary", "secret", "disclose"],
    "liability":       ["liabilit", "liable", "indemnif", "damages", "loss", "compensat"],
    "ip":              ["intellectual property", "copyright", "patent", "trademark", "trade secret", "invention", "ownership"],
    "dispute":         ["dispute", "arbitration", "mediation", "litigation", "jurisdiction", "court", "governing law"],
    "warranty":        ["warrant", "represent", "guarantee", "assur", "promise", "covenant"],
    "force_majeure":   ["force majeure", "act of god", "pandemic", "natural disaster", "beyond control"],
    "assignment":      ["assign", "transfer", "delegate", "novation", "successor"],
    "non_compete":     ["non-compet", "non compet", "restraint of trade", "exclusiv"],
}


def extract_clauses(text: str, clause_type: str = "all") -> dict:
    """
    Extract specific clause types from a contract.
    Returns dict of clause_type → list of relevant sentences.
    """
    sentences = re.split(r'(?<=[.!?])\s+|\n\n', text)
    sentences = [s.strip() for s in sentences if len(s.strip()) > 30]

    types_to_check = (
        list(CLAUSE_TYPES.keys())
        if clause_type == "all"
        else [clause_type] if clause_type in CLAUSE_TYPES else list(CLAUSE_TYPES.keys())
    )

    result = {}
    for ct in types_to_check:
        keywords = CLAUSE_TYPES[ct]
        found = []
        for sent in sentences:
            sent_lower = sent.lower()
            if any(kw in sent_lower for kw in keywords):
                if sent not in found:
                    found.append(sent[:300])
        if found:
            result[ct] = found[:6]  # max 6 per type

    return result


def format_clauses(clauses: dict, file_name: str = "") -> str:
    """Format extracted clauses as readable text."""
    lines = ["=" * 44, "CLAUSE EXTRACTION REPORT"]
    if file_name:
        lines.append(f"Source: {file_name}")
    lines.extend(["=" * 44, ""])

    if not clauses:
        lines.append("No specific clauses detected.")
        return "\n".join(lines)

    for clause_type, sentences in clauses.items():
        title = clause_type.replace("_", " ").upper()
        lines.append(f"── {title} ({len(sentences)} clauses)")
        for i, sent in enumerate(sentences, 1):
            lines.append(f"  {i}. {sent}")
        lines.append("")

    lines.append("=" * 44)
    return "\n".join(lines)


# ══════════════════════════════════════════════════════════════
# 3. DEADLINE & OBLIGATION TRACKER
# ══════════════════════════════════════════════════════════════

def track_deadlines(text: str, file_name: str = "") -> dict:
    """
    Extract all dates and obligations from a contract
    and return as a structured checklist.
    """
    dates       = _extract_dates(text)
    obligations = _extract_obligation_sentences(text)
    parties     = _extract_parties(text)

    # Try to pair dates with nearby obligation context
    date_items = []
    for date in dates:
        # Find sentence containing this date
        for sent in re.split(r'(?<=[.!?])\s+', text):
            if date in sent:
                context = sent.strip()[:200]
                date_items.append({
                    "date":    date,
                    "context": context,
                    "done":    False,
                })
                break
        else:
            date_items.append({"date": date, "context": "", "done": False})

    # Format obligations as checklist items
    obligation_items = []
    for obl in obligations[:10]:
        obligation_items.append({
            "text": obl[:200],
            "done": False,
        })

    return {
        "file_name":        file_name,
        "generated_at":     datetime.now().strftime("%Y-%m-%d %H:%M"),
        "parties":          parties,
        "total_dates":      len(date_items),
        "total_obligations": len(obligation_items),
        "dates":            date_items,
        "obligations":      obligation_items,
    }


def format_deadlines(tracker: dict) -> str:
    """Format deadline tracker as readable checklist."""
    lines = ["=" * 44, "DEADLINE & OBLIGATION TRACKER"]
    if tracker["file_name"]:
        lines.append(f"Source    : {tracker['file_name']}")
    lines.append(f"Generated : {tracker['generated_at']}")
    if tracker["parties"]:
        lines.append(f"Parties   : {', '.join(tracker['parties'][:2])}")
    lines.extend(["=" * 44, ""])

    lines.append(f"KEY DATES ({tracker['total_dates']} found)")
    lines.append("─" * 30)
    for item in tracker["dates"]:
        check = "☐"
        lines.append(f"{check}  {item['date']}")
        if item["context"]:
            lines.append(f"    Context: {item['context'][:120]}...")
        lines.append("")

    lines.append(f"OBLIGATIONS ({tracker['total_obligations']} found)")
    lines.append("─" * 30)
    for i, item in enumerate(tracker["obligations"], 1):
        lines.append(f"☐  {i}. {item['text']}")
        lines.append("")

    lines.append("=" * 44)
    lines.append("TIP: Print this checklist and tick off items as completed.")
    return "\n".join(lines)


# ══════════════════════════════════════════════════════════════
# 4. PROFESSIONAL REPORT GENERATOR (python-docx)
# ══════════════════════════════════════════════════════════════

def generate_report(
    extraction,
    file_name:   str = "",
    output_path: str = "",
    ai_analysis: str = "",
    comparison:  dict = None,
    clauses:     dict = None,
    tracker:     dict = None,
) -> str:
    """
    Generate a professional DOCX report using python-docx.
    Returns the output file path.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError("Run: pip install python-docx")

    doc = Document()

    # ── Page margins ──────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Styles helper ─────────────────────────────────────────
    def heading(text, level=1):
        p = doc.add_heading(text, level=level)
        p.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        return p

    def body(text):
        p = doc.add_paragraph(text)
        p.runs[0].font.size = Pt(11) if p.runs else None
        return p

    def bullet(text):
        doc.add_paragraph(text, style="List Bullet")

    def divider():
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.space_before = Pt(2)

    # ── Cover ─────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("VaultMind Legal Analysis Report")
    run.bold      = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x7C, 0x6A, 0xF7)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y at %H:%M')}")

    if file_name:
        fn = doc.add_paragraph()
        fn.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fn.add_run(f"Document: {file_name}").bold = True

    doc.add_paragraph()

    # ── Executive summary ─────────────────────────────────────
    heading("Executive Summary")
    body(f"Contract Type: {extraction.contract_type or 'General Contract'}")
    body(f"Governing Law: {extraction.governing_law or 'Not specified'}")
    body(f"Parties: {', '.join(extraction.parties) if extraction.parties else 'Not identified'}")
    body(f"Key Dates Found: {len(extraction.dates)}")
    body(f"Risk Clauses: {len(extraction.risk_clauses)}")
    body(f"Obligations: {len(extraction.obligations)}")
    divider()

    # ── Parties ───────────────────────────────────────────────
    if extraction.parties:
        heading("Parties", level=2)
        for p in extraction.parties:
            bullet(p)
        divider()

    # ── Key dates ─────────────────────────────────────────────
    if extraction.dates:
        heading("Key Dates & Deadlines", level=2)
        for d in extraction.dates:
            bullet(d)
        divider()

    # ── Risk clauses ──────────────────────────────────────────
    if extraction.risk_clauses:
        heading("Risk Clauses Detected", level=2)
        for r in extraction.risk_clauses:
            bullet(f"⚠  {r[:250]}")
        divider()

    # ── Obligations ───────────────────────────────────────────
    if extraction.obligations:
        heading("Key Obligations", level=2)
        for o in extraction.obligations:
            bullet(f"→  {o[:250]}")
        divider()

    # ── AI risk analysis ──────────────────────────────────────
    if ai_analysis:
        heading("AI Risk Analysis", level=2)
        for line in ai_analysis.split("\n"):
            line = line.strip()
            if line:
                body(line)
        divider()

    # ── Clause extraction ─────────────────────────────────────
    if clauses:
        heading("Clause Extraction", level=2)
        for clause_type, sentences in clauses.items():
            heading(clause_type.replace("_", " ").title(), level=3)
            for sent in sentences[:3]:
                bullet(sent[:250])
        divider()

    # ── Deadline tracker ──────────────────────────────────────
    if tracker:
        heading("Obligation Checklist", level=2)
        for item in tracker.get("obligations", [])[:8]:
            doc.add_paragraph(f"☐  {item['text'][:200]}", style="List Bullet")
        divider()

    # ── Comparison ────────────────────────────────────────────
    if comparison:
        heading("Contract Comparison", level=2)
        body(comparison.get("summary", ""))
        if comparison.get("risks_added"):
            heading("New Risks Added", level=3)
            for r in comparison["risks_added"]:
                bullet(f"⚠  {r[:200]}")
        divider()

    # ── Footer note ───────────────────────────────────────────
    doc.add_paragraph()
    note = doc.add_paragraph(
        "This report was generated by VaultMind — 100% local AI, no cloud processing. "
        "All analysis was performed on your machine. "
        "This report is for informational purposes only and does not constitute legal advice."
    )
    note.runs[0].font.size = Pt(9)
    note.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # ── Save ──────────────────────────────────────────────────
    if not output_path:
        stem = Path(file_name).stem if file_name else "document"
        output_path = str(
            Path("workspace/output") / f"{stem}_vaultmind_report.docx"
        )

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path

# ══════════════════════════════════════════════════════════════
# VaultMind v1.0 — 4 new features
# ══════════════════════════════════════════════════════════════
# ══════════════════════════════════════════════════════════════
# FEATURE 1: AMBIGUOUS LANGUAGE DETECTOR
# ══════════════════════════════════════════════════════════════

# Terms that are legally vague and routinely cause disputes
_AMBIGUOUS_TERMS = {
    "time": [
        "promptly", "immediately", "forthwith", "within a reasonable time",
        "reasonable time", "as soon as possible", "asap", "without delay",
        "timely", "expeditiously", "shortly",
    ],
    "quantity": [
        "reasonable", "adequate", "sufficient", "appropriate", "substantial",
        "significant", "material", "nominal", "fair", "proper", "satisfactory",
    ],
    "effort": [
        "best efforts", "best endeavours", "reasonable efforts",
        "reasonable endeavours", "commercially reasonable efforts",
        "good faith efforts", "all reasonable steps",
    ],
    "quality": [
        "good quality", "high quality", "acceptable quality",
        "industry standard", "professional standard", "workmanlike",
        "fit for purpose", "satisfactory",
    ],
    "scope": [
        "including but not limited to", "inter alia", "among other things",
        "and/or", "as applicable", "where applicable", "if applicable",
        "as necessary", "as required", "as appropriate",
    ],
    "obligation_strength": [
        "may", "might", "could", "should", "endeavour", "endeavor",
        "attempt", "try", "seek to", "use efforts",
    ],
    "monetary": [
        "nominal fee", "reasonable fee", "market rate", "prevailing rate",
        "fair value", "reasonable costs", "reasonable expenses",
        "arm's length", "arms length",
    ],
}

_SEVERITY = {
    "time": "HIGH",
    "quantity": "HIGH",
    "effort": "MEDIUM",
    "quality": "MEDIUM",
    "scope": "MEDIUM",
    "obligation_strength": "HIGH",
    "monetary": "HIGH",
}

_SUGGESTIONS = {
    "time": "Replace with a specific number of days/hours (e.g. 'within 7 business days').",
    "quantity": "Replace with a measurable threshold (e.g. 'not less than 90%' or a specific number).",
    "effort": "Specify 'best efforts' precisely or use a measurable KPI instead.",
    "quality": "Define the standard by reference to a specification, ISO standard, or objective test.",
    "scope": "List items explicitly rather than using open-ended qualifiers.",
    "obligation_strength": "'May' is permissive; 'shall' is mandatory. Confirm the intended obligation level.",
    "monetary": "Fix a specific amount or provide a clear formula for calculation.",
}


def detect_ambiguity(text: str, file_name: str = "") -> dict:
    """
    Scan a contract for legally vague / ambiguous language.
    Returns structured findings with severity and suggestions.
    No AI needed — instant, pattern-based.
    """
    sentences = re.split(r'(?<=[.!?])\s+|\n{2,}', text)
    sentences = [s.strip() for s in sentences if len(s.strip()) > 20]

    findings = []  # list of dicts
    seen_sentences = set()

    for sent in sentences:
        sent_lower = sent.lower()
        matched_categories = []

        for category, terms in _AMBIGUOUS_TERMS.items():
            for term in terms:
                if term in sent_lower:
                    matched_categories.append((category, term))
                    break  # one match per category per sentence is enough

        if matched_categories:
            # Deduplicate: don't report same sentence twice
            key = sent_lower[:80]
            if key in seen_sentences:
                continue
            seen_sentences.add(key)

            for category, term in matched_categories:
                findings.append({
                    "sentence": sent[:300],
                    "term": term,
                    "category": category,
                    "severity": _SEVERITY[category],
                    "suggestion": _SUGGESTIONS[category],
                })

    # Sort: HIGH first, then MEDIUM
    severity_order = {"HIGH": 0, "MEDIUM": 1, "LOW": 2}
    findings.sort(key=lambda x: severity_order.get(x["severity"], 2))

    high = sum(1 for f in findings if f["severity"] == "HIGH")
    medium = sum(1 for f in findings if f["severity"] == "MEDIUM")

    overall_risk = (
        "HIGH" if high >= 5 else
        "MEDIUM" if high >= 2 or medium >= 4 else
        "LOW"
    )

    return {
        "file_name": file_name,
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "total_findings": len(findings),
        "high_count": high,
        "medium_count": medium,
        "overall_risk": overall_risk,
        "findings": findings[:30],  # cap at 30 for readability
    }


def format_ambiguity(result: dict) -> str:
    """Format ambiguity detection result as readable text."""
    lines = [
        "=" * 50,
        "AMBIGUOUS LANGUAGE REPORT",
    ]
    if result["file_name"]:
        lines.append(f"Source   : {result['file_name']}")
    lines.append(f"Generated: {result['generated_at']}")
    lines.append("=" * 50)
    lines.append(f"Overall Risk  : {result['overall_risk']}")
    lines.append(f"Total Issues  : {result['total_findings']}")
    lines.append(f"  HIGH        : {result['high_count']}")
    lines.append(f"  MEDIUM      : {result['medium_count']}")
    lines.append("")

    if not result["findings"]:
        lines.append("✓ No significant ambiguous language detected.")
        lines.append("=" * 50)
        return "\n".join(lines)

    current_severity = None
    for i, f in enumerate(result["findings"], 1):
        if f["severity"] != current_severity:
            current_severity = f["severity"]
            lines.append(f"── {current_severity} SEVERITY")
            lines.append("─" * 40)

        lines.append(f"{i}. Term: \"{f['term']}\" [{f['category'].upper()}]")
        lines.append(f"   Clause: {f['sentence'][:200]}")
        lines.append(f"   Fix: {f['suggestion']}")
        lines.append("")

    lines.append("=" * 50)
    lines.append(
        "TIP: Vague terms are the #1 cause of contract disputes. "
        "Replace each flagged term with a specific, measurable standard."
    )
    return "\n".join(lines)


# ══════════════════════════════════════════════════════════════
# FEATURE 2: TIMELINE EXTRACTOR
# ══════════════════════════════════════════════════════════════

# Extend the date patterns already in legal_extractor with event context
_EVENT_KEYWORDS = [
    "effective", "commence", "begin", "start", "execute", "sign",
    "terminate", "expire", "end", "renew", "review", "deliver",
    "payment", "invoice", "due", "deadline", "notice", "comply",
    "submit", "complete", "achieve", "milestone", "close", "closing",
    "condition precedent", "longstop", "cut-off",
]

_DATE_PATTERNS_EXTENDED = [
    # ISO format
    r'\b(\d{4}[-/]\d{2}[-/]\d{2})\b',
    # dd Month YYYY / Month dd, YYYY
    r'\b(\d{1,2}(?:st|nd|rd|th)?\s+(?:January|February|March|April|May|June|'
    r'July|August|September|October|November|December)\s+\d{4})\b',
    r'\b((?:January|February|March|April|May|June|July|August|September|'
    r'October|November|December)\s+\d{1,2}(?:st|nd|rd|th)?,?\s+\d{4})\b',
    # Short month
    r'\b(\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)'
    r'\.?\s+\d{4})\b',
    # dd/mm/yyyy or mm/dd/yyyy
    r'\b(\d{1,2}[/-]\d{1,2}[/-]\d{4})\b',
    # Relative: within X days/months
    r'(within\s+\d+\s+(?:business\s+)?(?:days?|months?|weeks?|years?))',
    # Relative: X days/months after/from/before
    r'(\d+\s+(?:business\s+)?(?:days?|months?|weeks?)\s+'
    r'(?:after|from|before|following|prior to|of)(?:\s+the\s+\w+)?)',
    # Quarters / fiscal references
    r'(Q[1-4]\s+\d{4})',
    r'((?:first|second|third|fourth)\s+quarter(?:\s+of)?\s+\d{4})',
    # "on execution", "on signing"
    r'((?:on|upon)\s+(?:execution|signing|completion|delivery|closing))',
    # calendar / financial year
    r'((?:financial|fiscal|calendar)\s+year\s+\d{4})',
]

_COMPILED_DATE_RE = [re.compile(p, re.IGNORECASE) for p in _DATE_PATTERNS_EXTENDED]


def _find_dates_in_sentence(sentence: str):
    """Return list of date strings found in sentence."""
    found = []
    for pattern in _COMPILED_DATE_RE:
        for m in pattern.finditer(sentence):
            found.append(m.group(0).strip())
    return found


def _classify_event(sentence: str) -> str:
    """Classify what kind of event this timeline entry is."""
    s = sentence.lower()
    if any(w in s for w in ["effective", "commence", "begin", "start", "execution", "signing"]):
        return "Contract Start"
    if any(w in s for w in ["terminat", "expir", "end", "cessation"]):
        return "Termination / Expiry"
    if any(w in s for w in ["renew", "extension", "rollover"]):
        return "Renewal"
    if any(w in s for w in ["payment", "invoice", "fee", "remuneration", "salary"]):
        return "Payment"
    if any(w in s for w in ["deliver", "milestone", "complete", "submit", "achieve"]):
        return "Deliverable / Milestone"
    if any(w in s for w in ["notice", "notify", "written notice"]):
        return "Notice Requirement"
    if any(w in s for w in ["review", "audit", "assess"]):
        return "Review / Audit"
    if any(w in s for w in ["condition precedent", "longstop", "cut-off", "closing"]):
        return "Condition Precedent"
    return "General Obligation"


def extract_timeline(text: str, file_name: str = "") -> dict:
    """
    Extract a chronological timeline of events and obligations from a contract.
    Returns structured list of timeline entries sorted by type priority.
    """
    sentences = re.split(r'(?<=[.!?])\s+|\n{2,}', text)
    sentences = [s.strip() for s in sentences if len(s.strip()) > 20]

    entries = []
    seen = set()

    for sent in sentences:
        sent_lower = sent.lower()

        # Only process sentences that mention time/event keywords
        has_event_keyword = any(kw in sent_lower for kw in _EVENT_KEYWORDS)
        dates_in_sent = _find_dates_in_sentence(sent)

        if not (has_event_keyword or dates_in_sent):
            continue

        key = sent_lower[:80]
        if key in seen:
            continue
        seen.add(key)

        event_type = _classify_event(sent)
        entries.append({
            "event_type": event_type,
            "dates": dates_in_sent,
            "description": sent[:280],
            "has_specific_date": len(dates_in_sent) > 0,
        })

    # Sort: specific dates first, then by event type priority
    type_priority = {
        "Contract Start": 0,
        "Condition Precedent": 1,
        "Payment": 2,
        "Deliverable / Milestone": 3,
        "Notice Requirement": 4,
        "Renewal": 5,
        "Review / Audit": 6,
        "Termination / Expiry": 7,
        "General Obligation": 8,
    }
    entries.sort(key=lambda x: (
        0 if x["has_specific_date"] else 1,
        type_priority.get(x["event_type"], 9),
    ))

    return {
        "file_name": file_name,
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "total_entries": len(entries),
        "entries": entries[:25],  # top 25 most relevant
    }


def format_timeline(result: dict) -> str:
    """Format timeline as a readable chronological list."""
    lines = [
        "=" * 50,
        "CONTRACT TIMELINE",
    ]
    if result["file_name"]:
        lines.append(f"Source   : {result['file_name']}")
    lines.append(f"Generated: {result['generated_at']}")
    lines.append("=" * 50)
    lines.append(f"Total events found: {result['total_entries']}")
    lines.append("")

    if not result["entries"]:
        lines.append("No timeline events detected.")
        lines.append("=" * 50)
        return "\n".join(lines)

    current_type = None
    for i, entry in enumerate(result["entries"], 1):
        if entry["event_type"] != current_type:
            current_type = entry["event_type"]
            lines.append(f"── {current_type.upper()}")
            lines.append("─" * 40)

        date_str = ", ".join(entry["dates"]) if entry["dates"] else "No specific date"
        lines.append(f"{i}. [{date_str}]")
        lines.append(f"   {entry['description'][:200]}")
        lines.append("")

    lines.append("=" * 50)
    lines.append(
        "TIP: Cross-check this timeline with your calendar. "
        "Set reminders 7 days before each deadline."
    )
    return "\n".join(lines)